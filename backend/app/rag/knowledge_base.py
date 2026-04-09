"""RAG 知识库模块 - 基于Chroma向量数据库"""
import os
import hashlib
from typing import List, Dict, Any, Optional
from langchain_core.documents import Document
from app.core.llm import get_qwen_embedding
from app.core.multimodal import get_multimodal_processor
from app.config import get_settings
import chromadb
from chromadb.config import Settings as ChromaSettings

settings = get_settings()


class SimpleTextSplitter:
    """简单文本分词器"""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_text(self, text: str) -> List[str]:
        if not text or not text.strip():
            return []
        chunks = []
        start = 0
        text_len = len(text)
        while start < text_len:
            end = start + self.chunk_size
            chunk = text[start:end]
            if chunk:
                chunks.append(chunk)
            start = end - self.chunk_overlap
            if start >= text_len:
                break
        return chunks

    def split_documents(self, documents: List[Document]) -> List[Document]:
        result = []
        for doc in documents:
            if not doc.page_content or not doc.page_content.strip():
                continue
            chunks = self.split_text(doc.page_content)
            for i, chunk in enumerate(chunks):
                new_doc = Document(
                    page_content=chunk,
                    metadata={**doc.metadata, "chunk_index": i}
                )
                result.append(new_doc)
        return result


class KnowledgeBase:
    """RAG 知识库"""

    def __init__(self):
        self.embedding = get_qwen_embedding()
        self.collection_name = "psychology_knowledge"
        self.vector_store = None
        self.splitter = SimpleTextSplitter()
        self.multimodal_processor = get_multimodal_processor()
        self._initialize_vector_store()

    def _initialize_vector_store(self):
        db_path = settings.VECTOR_DB_PATH
        os.makedirs(db_path, exist_ok=True)

        self.vector_store = chromadb.PersistentClient(
            path=db_path,
            settings=ChromaSettings(
                anonymized_telemetry=False,
                allow_reset=True,
            )
        )

        try:
            self.collection = self.vector_store.get_collection(name=self.collection_name)
        except Exception:
            self.collection = self.vector_store.create_collection(
                name=self.collection_name,
                metadata={"description": "心理健康知识库"}
            )

    def _load_single_file(self, file_path: str) -> Optional[Document]:
        """加载单个文件"""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            content = ""

            if ext == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif ext == '.pdf':
                content = self._extract_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                content = self._extract_docx(file_path)
            elif ext == '.xlsx':
                content = self._extract_xlsx(file_path)
            elif ext == '.csv':
                content = self._extract_csv(file_path)
            elif ext == '.json':
                content = self._extract_json(file_path)
            elif ext == '.html' or ext == '.htm':
                content = self._extract_html(file_path)
            elif self.multimodal_processor.is_multimodal_file(file_path):
                # 处理多模态文件
                result = self.multimodal_processor.process_multimodal_file(file_path)
                if "error" in result:
                    return None
                content = f"""多模态内容
类型: {result['type']}
描述: {result['description']}
文件路径: {result['file_path']}
大小: {result.get('size', 0)} bytes
"""
            else:
                return None

            if not content.strip():
                return None

            metadata = {
                "source": file_path,
                "file_type": ext,
                "size": os.path.getsize(file_path),
                "last_modified": os.path.getmtime(file_path)
            }
            
            # 如果是多模态文件，添加额外元数据
            if self.multimodal_processor.is_multimodal_file(file_path):
                result = self.multimodal_processor.process_multimodal_file(file_path)
                if "error" not in result:
                    metadata["modality"] = result['type']
                    metadata["description"] = result['description']

            return Document(page_content=content, metadata=metadata)
        except Exception as e:
            print(f"Error loading file {file_path}: {str(e)}")
            return None

    def _extract_pdf(self, file_path: str) -> str:
        """提取PDF文本"""
        try:
            import fitz
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            return "[PDF解析需要安装pymupdf: pip install pymupdf]"
        except Exception as e:
            return f"[PDF解析失败: {str(e)}]"

    def _extract_docx(self, file_path: str) -> str:
        """提取Word文档文本"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        except ImportError:
            return "[Word解析需要安装python-docx: pip install python-docx]"
        except Exception as e:
            return f"[Word解析失败: {str(e)}]"

    def _extract_xlsx(self, file_path: str) -> str:
        """提取Excel文本"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            result = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                result.append(f"[Sheet: {sheet}]")
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        result.append(row_text)
            return "\n".join(result)
        except ImportError:
            return "[Excel解析需要安装openpyxl: pip install openpyxl]"
        except Exception as e:
            return f"[Excel解析失败: {str(e)}]"

    def _extract_csv(self, file_path: str) -> str:
        """提取CSV文本"""
        try:
            import csv
            result = []
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    result.append(" | ".join(row))
            return "\n".join(result)
        except Exception as e:
            return f"[CSV解析失败: {str(e)}]"

    def _extract_json(self, file_path: str) -> str:
        """提取JSON文本"""
        try:
            import json
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            return f"[JSON解析失败: {str(e)}]"

    def _extract_html(self, file_path: str) -> str:
        """提取HTML文本"""
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            return soup.get_text(separator='\n', strip=True)
        except ImportError:
            return "[HTML解析需要安装beautifulsoup4: pip install beautifulsoup4]"
        except Exception as e:
            return f"[HTML解析失败: {str(e)}]"

    def load_documents(self, knowledge_path: Optional[str] = None) -> List[Document]:
        """加载知识库文档"""
        if knowledge_path is None:
            knowledge_path = settings.KNOWLEDGE_BASE_PATH

        if not os.path.exists(knowledge_path):
            return []

        supported_exts = ('.md', '.txt', '.pdf', '.docx', '.doc', '.xlsx', '.csv', '.json', '.html')

        documents = []
        for root, dirs, files in os.walk(knowledge_path):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                file_path = os.path.join(root, file)
                # 检查是否为支持的文件类型或多模态文件
                if ext in supported_exts or self.multimodal_processor.is_multimodal_file(file_path):
                    doc = self._load_single_file(file_path)
                    if doc:
                        documents.append(doc)

        return self.splitter.split_documents(documents)

    def add_documents(self, documents: List[Document], category: str = "general") -> Dict[str, Any]:
        """添加文档到知识库"""
        if not documents:
            return {"success": False, "error": "文档为空"}

        try:
            texts = [doc.page_content for doc in documents]
            metadatas = []
            for doc in documents:
                meta = dict(doc.metadata)
                meta["category"] = category
                metadatas.append(meta)

            ids = []
            for i, text in enumerate(texts):
                hash_id = hashlib.md5(text.encode()).hexdigest()[:16]
                ids.append(f"doc_{category}_{hash_id}_{i}")

            embeddings = self.embedding.embed_documents(texts)

            self.collection.add(
                embeddings=embeddings,
                documents=texts,
                metadatas=metadatas,
                ids=ids,
            )

            return {
                "success": True,
                "count": len(documents),
                "ids": ids[:5]
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def add_text(self, text: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """添加单条文本到知识库"""
        try:
            if metadata is None:
                metadata = {}

            hash_id = hashlib.md5(text.encode()).hexdigest()[:16]
            doc_id = f"doc_single_{hash_id}"

            embedding = self.embedding.embed_query(text)

            self.collection.add(
                embeddings=[embedding],
                documents=[text],
                metadatas=[metadata],
                ids=[doc_id],
            )

            return {"success": True, "id": doc_id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def search(self, query: str, top_k: int = 3, category: str = None, score_threshold: float = 0.3) -> List[Dict[str, Any]]:
        """检索最相关的文档"""
        try:
            query_embedding = self.embedding.embed_query(query)

            where_filter = {"category": category} if category else None

            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                where=where_filter
            )

            if not results or not results.get("documents"):
                return []

            documents = results["documents"][0]
            metadatas = results.get("metadatas", [[]])[0]
            distances = results.get("distances", [[]])[0]
            ids = results.get("ids", [[]])[0]

            # 过滤低相似度结果并排序
            filtered_results = []
            for doc, meta, dist, doc_id in zip(documents, metadatas, distances, ids):
                score = round(1 - dist, 4) if dist else 0
                if score >= score_threshold:
                    filtered_results.append({
                        "id": doc_id,
                        "content": doc,
                        "metadata": meta,
                        "score": score,
                    })

            # 按相似度分数排序
            filtered_results.sort(key=lambda x: x['score'], reverse=True)
            return filtered_results
        except Exception as e:
            return [{"error": str(e), "content": "检索失败"}]

    def get_all_categories(self) -> List[str]:
        """获取所有分类"""
        try:
            results = self.collection.get()
            if not results or not results.get("metadatas"):
                return []

            categories = set()
            for meta in results["metadatas"]:
                if meta and "category" in meta:
                    categories.add(meta["category"])

            return sorted(list(categories))
        except Exception:
            return []

    def get_collection_stats(self) -> Dict[str, Any]:
        """获取知识库统计信息"""
        try:
            count = self.collection.count()
            categories = self.get_all_categories()

            return {
                "success": True,
                "total_documents": count,
                "categories": categories,
                "category_count": len(categories)
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def delete_by_category(self, category: str) -> Dict[str, Any]:
        """删除指定分类的文档"""
        try:
            results = self.collection.get(where={"category": category})
            if results and results.get("ids"):
                self.collection.delete(ids=results["ids"])
                return {"success": True, "deleted_count": len(results["ids"])}
            return {"success": True, "deleted_count": 0}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def reset(self):
        """重置知识库"""
        try:
            self.vector_store.delete_collection(name=self.collection_name)
            self._initialize_vector_store()
        except Exception:
            pass

    def get_all_documents(self, limit: int = 100) -> List[Dict[str, Any]]:
        """获取所有文档"""
        try:
            results = self.collection.get(limit=limit)

            if not results or not results.get("documents"):
                return []

            return [
                {
                    "id": doc_id,
                    "content": doc,
                    "metadata": meta,
                }
                for doc_id, doc, meta in zip(results["ids"], results["documents"], results.get("metadatas", [[]]))
            ]
        except Exception as e:
            return [{"error": str(e)}]


_knowledge_base_instance = None


def get_knowledge_base() -> KnowledgeBase:
    """获取知识库单例"""
    global _knowledge_base_instance
    if _knowledge_base_instance is None:
        _knowledge_base_instance = KnowledgeBase()
    return _knowledge_base_instance


def initialize_knowledge_base() -> Dict[str, Any]:
    """初始化知识库"""
    try:
        kb = get_knowledge_base()
        documents = kb.load_documents()

        if not documents:
            return {"success": True, "document_count": 0}

        result = kb.add_documents(documents)
        return {
            "success": True,
            "document_count": len(documents),
            "result": result
        }
    except Exception as e:
        return {"success": False, "error": str(e)}